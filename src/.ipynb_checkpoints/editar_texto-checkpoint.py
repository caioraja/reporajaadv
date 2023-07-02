from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# TODO configurar a edição do texto parei em p107
def editar_texto(doc, dic):
    """_summary_

    Args:
        dic (dict): Recebe um dicionário com todas as respostas do questionário e promove as alterações no arquivo docx
    """
    for par in doc.paragraphs:
        run = par.runs[0]
        run.doc.styles["Normal"].font
        run.name = "Arial"
        run.size = Pt(12)
    #font = doc.styles["Normal"].font
    #font.name = "Arial"
    #font.size = Pt(12)

    MARCA_DO_CLIENTE = dic["historia"][0]
    
    p3 = doc.paragraphs[3]
    p3.clear()
    p3.add_run("01.").bold = True
    p3.add_run(
        "\tDeclaro que, nesta data, recebi um exemplar da Circular de Oferta de Franquia (COF), com todas as informações relevantes sobre o Sistema de Franquias da Rede de Franquias "
    )
    p3.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p3.add_run(
        ", exigidas pela Lei nº 13.966/2019 (Lei de Franquias Empresariais), afirmando minha ciência de que, dentro do período de 10 (dez) dias seguintes ao dia de hoje, não assinarei nenhum Contrato ou Pré-Contrato de Franquia, e, tampouco, promoverei o pagamento de qualquer quantia para a Franqueadora ou pessoa a ela ligada."
    )

    p9 = doc.paragraphs[9]
    p9.clear()
    p9.add_run("04.").bold = True
    p9.add_run(
        "\tDeclaro minha ciência quanto ao fato de que a entrega desta COF não representa certeza, e, tampouco, expectativa da minha aprovação no processo de avaliação de candidatos a franqueados da "
    )
    p9.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p9.add_run(
        ", ficando claro que a Franqueadora não está obrigada a externar suas motivações para minha eventual reprovação em tal procedimento."
    )

    p11 = doc.paragraphs[11]
    p11.clear()
    p11.add_run("05.").bold = True
    p11.add_run(
        "\tDeclaro que me obrigo a não reproduzir a documentação que me está sendo entregue, me obrigo a preservar o sigilo e confidencialidade das informações que a mim estão sendo compartilhadas, e me comprometo a não utilizar as informações compartilhadas pela Franqueadora, para estruturar e/ou operar outro negócio concorrente à Rede de Franquias Empresariais "
    )
    p11.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p11.add_run(
        ", sob pena de sujeitar-me à responsabilização cível e penal pela prática de concorrência desleal, nos termos da Lei nº 9.279/1996 (Lei de Propriedade Industrial) e demais normativos pertinentes, com pagamento de indenização por perdas e danos arbitrada no valor mínimo de R$ 100.000,00 (cem mil reais)."
    )

    p13 = doc.paragraphs[13]
    p13.clear()
    p13.add_run("06.").bold = True
    p13.add_run(
        "\tDeclaro que autorizo o tratamento dos meus dados pessoais, pela Franqueadora do Sistema de Franquias Empresariais "
    )
    p13.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p13.add_run(
        " e empresas a ela ligadas, ciente de que isso se fará necessário e imprescindível para a adoção de procedimentos preliminares relacionados à minha avaliação e eventual eleição ao posto de Franqueado da Rede – ciente de que meus dados pessoais serão utilizados para realizar pesquisas em órgãos públicos e/ou de avaliação de crédito, obter certidões e, se o caso, preencher o Contrato de Franquia que eventualmente poderá firmado no futuro –, ainda que a celebração de Contrato de Franquia com a Franqueadora venha a não ser concretizada, observando-se os termos do artigo 7º, incisos I e V da Lei nº 13.709/2018) (Lei Geral de Proteção de Dados), dentre outros. Neste ato, também me declaro ciente de que, em qualquer momento, poderei pleitear a interrupção do tratamento dos meus dados pessoais e, se o caso, a sua eliminação, interrompendo o prosseguimento dos procedimentos preliminares acima mencionados."
    )

    p66 = doc.paragraphs[66]
    p66.clear()
    p66.add_run(
        "Você está recebendo nossa Circular de Oferta de Franquia (COF), contendo informações e dados básicos sobre quem somos, sobre nossa filosofia de trabalho, nossos anseios, e acima de tudo, sobre o modelo de franchising do negócio "
    )
    p66.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p66.add_run(" – de agora em diante chamado apenas ")
    p66.add_run(f'"{MARCA_DO_CLIENTE}"').bold = True
    p66.add_run(".")

    p70 = doc.paragraphs[70]
    p70.clear()
    p70.add_run(
        "Observa-se que a COF não registra, por si só, qualquer obrigação ou compromisso, servindo apenas como uma fonte objetiva de informações a respeito da Franqueadora, seu negócio e a sua operação franqueada, auxiliando você a tomar sua decisão de ingresso ou não ingresso na Rede de Franquias "
    )
    p70.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p70.add_run(".")

    p78 = doc.paragraphs[78]
    p78.clear()
    p78.add_run("Time da ")
    p78.add_run(f"{MARCA_DO_CLIENTE}").bold = True

    p100 = doc.paragraphs[100]
    p100.clear()
    p100.add_run("O que motivou a criar o projeto ")
    p100.add_run(f"{MARCA_DO_CLIENTE} ").bold = True
    p100.add_run("?")

    p101 = doc.paragraphs[101]
    p101.clear()
    p101.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p101.add_run(f"{dic['historia'][1]} ").italic = True

    p103 = doc.paragraphs[103]
    p103.clear()
    p103.add_run("*Relato inspirador dos fundadores da operação ")
    p103.add_run(f"{MARCA_DO_CLIENTE} ").bold = True

    p105 = doc.paragraphs[105]
    p105.clear()
    p105.add_run("Ao que a ")
    p105.add_run(f"{MARCA_DO_CLIENTE} ").bold = True
    p105.add_run(" se propõe?")

    p107 = doc.paragraphs[107]
    p107.clear()
    p107.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p107.add_run(f"{dic['historia'][2]} ")

    p109 = doc.paragraphs[109]
    p109.clear()
    p109.add_run(
        "Depois de muita dedicação e incansável perseverança dos seus fundadores e colaboradores, a "
    )
    p109.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    if dic["COF"] == 1:
        p109.add_run(
            f" se estabeleceu como referência regional no segmento de {dic['modelo'][2].lower()}, conquistando clientes em todo mercado brasileiro. "
        )
    else:
        p109.add_run(
            f" se estabeleceu como referência regional no segmento de comercialização de produtos de {dic['modelo'][2].lower()}, conquistando clientes em todo mercado brasileiro. "
        )

    p111 = doc.paragraphs[111]
    p111.clear()
    p111.add_run("O que a ")
    p111.add_run(f"{MARCA_DO_CLIENTE} ").bold = True
    p111.add_run(
        " orgulhosamente fez, preservando sua visão, sua missão e seus valores:"
    )

    p113 = doc.paragraphs[113]
    p113.clear()
    p113.add_run("Visão: ")
    p113.add_run(f"{dic['historia'][4]}").italic = True

    p115 = doc.paragraphs[115]
    p115.clear()
    p115.add_run("Missão: ")
    p115.add_run(f"{dic['historia'][5]}").italic = True

    p117 = doc.paragraphs[117]
    p117.clear()
    p117.add_run("Valores: ")
    p117.add_run(f"{dic['historia'][6]}").italic = True

    p119 = doc.paragraphs[119]
    p119.clear()
    p119.add_run(
        "Ao sistematizar seu negócio sob um formato de franquia empresarial (franchising), com início em 2022, o time da "
    )
    p119.add_run(f"{MARCA_DO_CLIENTE} ").bold = True
    p119.add_run(
        " se propõe a auxiliar pessoas, que apresentem identidade com a sua marca e o seu negócio, a tentarem operar em tal especial mercado, à frente de um ponto "
    )
    if dic["COF"] == 1:
        p119.add_run(f"prestação de serviços de {dic['modelo'][2].lower()}")
    else:
        p119.add_run(f"comercialização de produtos de {dic['modelo'][2].lower()}")
    p119.add_run(", sob a bandeira ")
    p119.add_run(
        ", na modalidade de franquia escolhida, utilizando o know-how transmitido pela Franqueadora, de acordo com esta COF, o Contrato de Franquia Empresarial oportunamente firmado, e demais instrumentos que eventualmente venham a compor a formatação jurídica do Sistema de Franquias "
    )
    p119.add_run(f"{MARCA_DO_CLIENTE}.").bold = True

    if dic["modelo"][0] != 1:
        p121 = doc.paragraphs[121]
        p121.clear()
        p121.add_run("Atualmente, a ")
        p121.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p121.add_run(
            f" estruturou  0{dic['modelo'][0]} ({dic['modelo'][1]}) modelos operacionais distintos, assim identificados: “M.O. 01”; “M.O. 02”; e “M.O. 03”; que apresentam como principais diferenças:"
        )
        p131 = doc.paragraphs[131]
        p131.add_run(
            "05. Dentre outras distinções tratadas ao longo desta COF e durante o processo de apresentação do Sistema de Franquias "
        )
        p133 = doc.paragraphs[133]
        p133.clear()

    else:
        for i in [121, 123, 125, 127, 129, 131]:
            par = doc.paragraphs[i]
            par.clear()

        p133 = doc.paragraphs[133]
        p133.clear()
        p133.add_run("Atualmente, ")
        p133.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p133.add_run(
            " estruturou 01 (um) modelo operacional identificado como “M.O. 01”, cujas principais características e informações iniciais serão tratadas ao longo desta COF e por meio de outros documentos disponibilizados durante o processo de apresentação do Sistema de Franquias "
        )
        p133.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p133.add_run(".")

    p135 = doc.paragraphs[135]
    p135.clear()
    p135.add_run("Esta COF é um dos instrumentos que integram a formatação jurídica do Sistema de Franquias Empresariais (")
    p135.add_run("franchising").italic = True
    p135.add_run(") da Rede de Franquias ")
    p135.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p135.add_run(", contendo informações relevantes sobre o negócio em destaque, e sobre os termos do relacionamento pretendido entre Franqueadora e Franqueados.")

    p139 = doc.paragraphs[139]
    p139.clear()
    p139.add_run("Você, candidato que está iniciando o processo de seleção para integrar uma das unidades da Rede de Franquias ")
    p139.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p139.add_run(", sinta-se à vontade para, durante ou após a leitura desta COF, solicitar qualquer esclarecimento adicional.")


    p151 = doc.paragraphs[151]
    p151.clear()
    p151.add_run("Obviamente, ")
    p151.add_run("mesmo em um Sistema de Franchising, os riscos inerentes a quaisquer atividades empresariais ainda existem").bold = True
    p151.add_run(". De modo que nenhuma Franquia Empresarial promete êxito aos seus franqueados, e nem a Franqueadora e a Rede de Franquias ")
    p151.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p151.add_run(" poderá fazê-lo.")

    p153 = doc.paragraphs[153]
    p153.clear()
    p153.add_run("Frisa-se, pois, desde já, que a Rede de Franquias ").bold = True
    p153.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p153.add_run(" não está a prometer aos seus franqueados enriquecimento, lucro ou êxito").bold = True
    p153.add_run(", haja vista que o atingimento de tais fins depende não só do apoio da Franqueadora, mas, principalmente, da dedicação efetiva do Franqueado, para o desenvolvimento do negócio, além de contar com um cenário econômico favorável, bom desempenho do setor etc.")

    p155 = doc.paragraphs[155]
    p155.clear()
    p155.add_run("Por qual(ais) razão(ões), afinal, pode se mostrar interessante investir em uma operação do segmento de ")
    p155.add_run(f"\"{dic['modelo'][2]}\"").italic = True
    p155.add_run(", desenvolvida em um sistema de franquia empresarial – ")
    p155.add_run("franchising").italic = True
    p155.add_run("?")

    p157 = doc.paragraphs[157]
    p157.clear()
    p157.add_run("Mesmo em um cenário de crise econômica como aquele recentemente enfrentado, considerando números disponíveis em novembro de 2022, tem-se que ")
    p157.add_run("o franchising").italic = True
    p157.add_run(" teve aumento de mais de 12,6% no faturamento acumulado entre o quarto trimestre de 2021 e o quarto trimestre de 2022, sendo considerado a melhor performance dos últimos três anos no período (T4), indicando que o setor não apenasse recuperou plenamente do período pandêmico como mostra claros sinais de crescimento. ")
    
    p159 = doc.paragraphs[159]
    p159.clear()
    p159.add_run(f"O segmento de \"{dic['modelo'][2]}\" – no qual está inserida a atividade principal desenvolvida dentro do Sistema de Franquias da Rede ")
    p159.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p159.add_run(f" – considerando números referentes ao terceiro trimestre de 2022, apresentou uma variação positiva de faturamento de aproximadamente {dic['modelo'][3]}%, além de uma variação de {dic['modelo'][4]}% na criação de novas unidades, considerando o mesmo período no ano anterior.²")

    p161 = doc.paragraphs[161]
    p161.clear()
    p161.add_run("Em tais linhas gerais, apresenta-se o potencial de crescimento e consolidação da atividade desempenhada no ambiente da Rede de Franquias ")
    p161.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p161.add_run(', dentro do ')
    p161.add_run("franchising").italic = True

    
    #TODO RESOLVER PROBLEMA DA FONTE
    p163 = doc.paragraphs[163]
    p163.clear()
    p163.add_run("Como dito acima, o Sistema de Franquias ")
    p163.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p163.add_run("CLIENTE foi estruturado em 2022, iniciando sua operação logo após, com o escopo de auxílio de pessoas, que apresentem identidade com a sua marca e o seu negócio, a tentarem operar no seu mercado, à frente de um ponto de ")
    if dic["COF"] == 1:
        p163.add_run(f"comercialização e prestação de serviços de {dic['modelo'][2].lower()}")
    else:
        p163.add_run(f"comercialização de produtos de {dic['modelo'][2].lower()}")
    p163.add_run(" – sendo este o chamado “negócio franqueado” –, sob a bandeira ")
    p163.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p163.add_run(", na modalidade de franquia escolhida, utilizando o know-how transmitido pela Franqueadora, de acordo com esta COF demais instrumentos que eventualmente venham a compor a formatação do Sistema de Franquias ")
    p163.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p163.add_run(".")

    p165 = doc.paragraphs[165]
    p165.clear()
    p165.add_run("Especialmente no que diz respeito às atividades que se espera que sejam desempenhadas pelos Franqueados da Rede de Franquias ")
    p165.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p165.add_run(", pode-se dizer, basicamente, que eles devem:")

    if dic['modelo'][5] == 1:
        p189 = doc.paragraphs[189]
        p189.clear()
        p189.add_run("Especificamente nas operações montadas tendo como inspiração a modalidade operacional “home based”, se dispensa a necessidade de observância de layout e fachada, mantendo-se as demais requisições da Franqueadora.”")

        p191 = doc.paragraphs[191]
        p191.clear()
        p191.add_run("Os Franqueados deverão respeitar, rigorosamente, os padrões de uso da marca ")
        p191.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p191.add_run(", a tratando de acordo com todas as especificações transmitidas pela Franqueadora, em iniciativas de publicidade, uniformes de funcionários – caso operem com funcionários –, identidade visual da sua operação em geral, com o intuito de reforçar a identidade da ")
        p191.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p191.add_run(", perante todo mercado.")
    else:
        p189 = doc.paragraphs[189]
        p189.clear()
        p189.add_run("A observância dos padrões estabelecidos para montagem da unidade franqueada, é imprescindível: os Franqueados deverão utilizar, rigorosamente, layout, fachada e cores indicadas pela Franqueadora, e, ainda o padrão dos uniformes dos funcionários que trabalharão na unidade franqueada. Deverão, em suma tratar a marca MARCA DO CLIENTE de acordo com todas as especificações transmitidas pela Franqueadora, com o intuito de reforçar a identidade da ")
        p189.add_run(f"{MARCA_DO_CLIENTE}").bold = True
        p189.add_run(" CLIENTE de acordo com todas as especificações transmitidas pela Franqueadora, com o intuito de reforçar a identidade da ")    
        p189.add_run(", perante todo mercado.")

        p191 = doc.paragraphs[191]
        p191.clear()
    
    p193 = doc.paragraphs[193]
    p193.clear()
    p193.add_run("Observa-se que os Franqueados da Rede ")
    p193.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p193.add_run(", devem atender, outrossim, as regras dos órgãos regulamentadores das suas atividades. ")
    
    p195 = doc.paragraphs[195]
    p195.clear()
    p195.add_run("Os insumos para prestação de serviços (e eventuais produtos relacionados) comercializados dentro das operações ")
    p195.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p195.add_run(", somente poderão ser adquiridos junto à Franqueadora ou juntamente aos seus fornecedores homologados.")

    p205 = doc.paragraphs[205]
    p205.clear()
    p205.add_run("A autorização de uso da marca ")
    p205.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p205.add_run(" será concedida pela Franqueadora, para os Franqueados, a título não exclusivo, precário, para utilização apenas na(s) sua(s) própria(s) unidade(s) franqueada(s).")

    p209 = doc.paragraphs[209]
    p209.clear()
    p209.add_run("Caso obtenha êxito na assunção da condição de Franqueado da Rede ")
    p209.add_run(f"{MARCA_DO_CLIENTE}").bold = True
    p209.add_run(", esteja certo de que você contará com todo apoio possível da Franqueadora, para tentar se estabelecer em tal cenário. ")



    doc.save("Teste.docx")
