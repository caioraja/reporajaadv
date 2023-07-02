from docx.enum.text import WD_ALIGN_PARAGRAPH

def change_marca_cliente(doc, MARCA_DO_CLIENTE):
    for run in doc.paragraphs:
        for par in run.runs:
            if par.text == "MARCA DO CLIENTE":
                par.text = MARCA_DO_CLIENTE

def change_modelo(doc, palavra, modelo):
    for run in doc.paragraphs:
        for par in run.runs:
            if par.text == palavra:
                par.text = modelo
                
def add_texto(doc, n_paragrafo, n_run, texto):
    run = doc.paragraphs[n_paragrafo]
    run.runs[n_run].text = f" {texto}" if n_run != 0 else texto
    run.aligment = WD_ALIGN_PARAGRAPH.JUSTIFY

def funcionario_ideal(doc, dic, par, index):
    for run in doc.paragraphs[par]:
        for par in run.runs:
            if par.text == "AAAAA":
                par.text = dic['Funcionario Ideal'][index]

def ed_texto(doc, dic):
    add_texto(doc, 101, 0, dic["historia"][1])
    add_texto(doc, 107, 0, dic["historia"][2])
    add_texto(doc, 113, 1, dic["historia"][4])
    add_texto(doc, 115, 1, dic["historia"][5])
    add_texto(doc, 117, 1, dic["historia"][6])

    run109 = doc.paragraphs[109]
    run109.runs[
        7
    ].text = f"comercialização e prestação de serviços de {dic['historia'][3]}"
    run109.runs[8].clear()
    run109.runs[9].text = ","

    run119 = doc.paragraphs[119]
    run119.runs[
        23
    ].text = f"comercialização e prestação de serviços de {dic['historia'][3]}"
    run119.runs[24].clear()
    run119.runs[25].clear()

    if dic["n_Modelos"][0] == 1:
        for i in range(121, 133, 2):
            runi = doc.paragraphs[i]
            runi.clear()
    else:
        run133 = doc.paragraphs[133]
        run133.clear()

    run155 = doc.paragraphs[155]
    run155.runs[1].text = dic["Segmento"][0]
    run159 = doc.paragraphs[159]
    run159.runs[1].text = dic["Segmento"][0]
    run159.runs[12].text = dic["Segmento"][1]
    run163 = doc.paragraphs[163]
    run163.runs[
        18
    ].text = f"comercialização e prestação de serviços de {dic['historia'][3]}"
    run163.runs[19].clear()
    run163.runs[20].clear()

    if dic["Segmento"][2] != 1:
        run189 = doc.paragraphs[189]
        for i in run189.runs[7:]:
            i.clear()

    run221 = doc.paragraphs[221]
    run221.runs[
        4
    ].text = f"sendo obrigatório aos Franqueados o adequado preenchimento dos {dic['Segmento'][3]}"

    texto = "".join(
        f'Treinamento de {dic["Treino_franqueadora"][i]}; '
        for i in range(len(dic["Treino_franqueadora"]) - 1)
        if dic["Treino_franqueadora"][i] != 0
    )
    if dic["Treino_franqueadora"][-1] == 0:
        texto = f"{texto[:-2]}."
    else:
        texto = f"{texto[:-2]} e Treinamento {dic['Treino_franqueadora'][-1]}."
    run237 = doc.paragraphs[237]
    run237.runs[2].text = texto
    for i in run237.runs[4:]:
        i.clear()

    if dic["Treinamento"][0] == 1 and dic["Treinamento"][3] == 3:
        for i in [247, 249, 251]:
            runx = doc.paragraphs[i]
            runx.clear()

    elif dic["Treinamento"][0] == 1 and dic["Treinamento"][3] == 1:
        for i in [245, 249, 251]:
            runx = doc.paragraphs[i]
            runx.clear()

    elif dic["Treinamento"][0] == 0 and dic["Treinamento"][3] == 3:
        for i in [245, 247, 251]:
            runx = doc.paragraphs[i]
            runx.clear()
    else:
        for i in [245, 247, 249]:
            runx = doc.paragraphs[i]
            runx.clear()

    run255 = doc.paragraphs[255]
    run255.runs[1].text = f"0{dic['Treinamento'][1]} ({dic['Treinamento'][2]}) dias"

    if dic["Treinamento"][4] == 1:
        run259 = doc.paragraphs[259]
        run259.clear()
    else:
        run257 = doc.paragraphs[257]
        run257.clear()

    texto = "".join(
        f'Manual de {dic["Manuais"][i]}; '
        for i in range(len(dic["Manuais"]) - 1)
        if dic["Manuais"][i] != 0
    )
    if dic["Manuais"][-1] == 0:
        texto = f"{texto[:-2]}."
    else:
        texto = f"{texto[:-2]} e Manual de {dic['Manuais'][-1]}."
    run271 = doc.paragraphs[271]
    run271.runs[-1] = texto

    if dic["Segmento"] == 1:
        run277 = doc.paragraphs[277]
        run277.clear()
    else:
        for x in [278, 281, 283]:
            runx = doc.paragraphs[x]
            runx.clear()

    texto = ", incluindo-se " + "".join(
        f' {dic["Layout"][i]}; '
        for i in range(len(dic["Layout"]) - 1)
        if dic["Layout"][i] != 0
    )
    if dic["Layout"][-1] == 0:
        texto = f"{texto[:-2]}."
    else:
        texto = f"{texto[:-2]} e {dic['Layout'][-1]}."
    run280 = doc.paragraphs[280]
    run280.runs[11].text = texto

    run285 = doc.paragraphs[285]
    if dic["Produtos"][0] != 0:
        run285.runs[1].text = f"prestando serviços de  + {dic['Produtos'][0]}"
    else:
        run285.clear()


    doc.save("T2.docx")
